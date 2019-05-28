from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import datetime


# This class takes a template file and then fills in the correct data. It then will
# get emailed for approval.
# To instance form gen one should provide the correct OC insturctor name, OC department,
# military_course, and the oc_course
class FileGen:

    # initializing variables to be used throught the generator
    def __init__(self, instructor_name, department, military_course, oc_course):
        self.doc = Document('Test.docx')  # test.docx should be the name of the template file.
        self.total_tables = self.doc.tables  # locating document tables
        self.comp_table = self.total_tables[0]  # locating correct table for course comparason
        self.date_cell = self.comp_table.cell(0, 0)
        self.military_course_cell = self.comp_table.cell(0, 1)
        self.instructor_name_cell = self.comp_table.cell(1, 0)  # data for filling in top of doc.
        self.oc_course_cell = self.comp_table.cell(1, 1)
        self.instr_name = instructor_name
        self.dep_name = department
        self.m_course = military_course
        self.oc_course = oc_course
        # bellow marks out the columns for both military and olivet course outcomes.
        self.mc_row = 3
        self.mc_column = 1
        self.oc_row = 3
        self.oc_column = 0
        self.__Fill_Course_Info()

    # will add checkboxes the the correct columns
    def __Add_Checkbox(self, jst_outcome):
        column_check_add = self.mc_column + 1

        for i in range(column_check_add, len(self.comp_table.columns)):
            cell_check_add = self.comp_table.cell(self.mc_row, i)
            para = cell_check_add.paragraphs[0]
            para.text += "\u2610"
            para_format = para.paragraph_format
            para_format.space_before = Pt(1)
            if len(jst_outcome) > 50:
                para.text += "\n\n\n"
            else:
                para.text += "\n\n"
            para_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # called during initualization fills in the top of the table.
    def __Fill_Course_Info(self):
        now = datetime.datetime.now()
        self.date_cell.text = "Date of Initiation:\n" + str(now.month) + "-" + str(now.day) + "-" + str(now.year)
        self.military_course_cell.text = "JST or AU course for which Credit/Equivalency is sought:\n" + self.m_course
        self.instructor_name_cell.text = "Evaluator Name:\n" + self.instr_name + "\nDepartment:\n" + self.dep_name
        self.oc_course_cell.text = "Olivet College course being considered for possible equivalency:\n" + self.oc_course

    # used for entering individual Olivet course outcomes
    def Olivet_Course_Outcomes(self, c_outcome):
        self.oc_outcome_cell = self.comp_table.cell(self.oc_row, self.oc_column)
        self.oc_outcome_cell.text = c_outcome
        self.oc_row += 1

    # used for entering individual Military course outcomes
    def JST_Outcomes(self, jst_outcome, new_cell=False):  # used when the user wants to move to the next cell.
        if new_cell == True:
            self.mc_row += 1
        self.mc_outcome_cell = self.comp_table.cell(self.mc_row, self.mc_column)
        self.mc_outcome_cell.text += jst_outcome + "\n\n"
        self.__Add_Checkbox(jst_outcome)

    # adds both the Olivet course outcomes with there coresponding military course outcomes.
    # c_outcomes would be just a string for the Olivet college outcome and then jst_outcomes would be the coresponding array of matching military outcomes.
    def Like_Outcomes(self, c_outcome, jst_outcome):
        self.Olivet_Course_Outcomes(c_outcome)
        for i in range(0, len(jst_outcome)):
            self.JST_Outcomes(jst_outcome[i])
        self.mc_row += 1

    # will be used if we decide to implement emailing the form.
    def Email_Doc(self):
        pass

    # used to save the document.
    def Save_Doc(self):
        self.doc.save('Test-Saved.docx')
